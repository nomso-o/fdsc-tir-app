from typing import List, Dict
import io

from docx import Document
import fitz  # PyMuPDF


def build_docx_from_results(results: List[Dict]) -> bytes:
    doc = Document()
    for r in results:
        doc.add_heading(f"TIR: {r['tir_blob_path']}", level=1)
        doc.add_paragraph(r["rationale"])
        doc.add_paragraph(r["markdown_table"])
        doc.add_page_break()
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()


def build_pdf_from_results(results: List[Dict]) -> bytes:
    pdf_doc = fitz.open()
    for r in results:
        page = pdf_doc.new_page()
        text = f"TIR: {r['tir_blob_path']}\n\n{r['rationale']}\n\n{r['markdown_table']}"
        page.insert_text((50, 50), text, fontsize=10)
        # Note: for long text, you'd want to paginate/flow text properly.
    data = pdf_doc.tobytes()
    pdf_doc.close()
    return data
